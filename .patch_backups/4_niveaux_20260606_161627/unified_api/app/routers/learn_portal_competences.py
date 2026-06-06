from fastapi import APIRouter, HTTPException, Request, Response, UploadFile, File, Form
from pydantic import BaseModel
from typing import Optional, Any
from psycopg.rows import dict_row
from psycopg.types.json import Json
from io import BytesIO
import uuid
import os
import json
import re
import unicodedata

from app.routers.skills_portal_common import get_conn
from app.routers.learn_portal_common import learn_require_user, learn_fetch_profile
from app.routers.skills_portal_pdf_common import (
    build_pdf_document,
    build_competence_pdf_story,
)

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

router = APIRouter()


# ------------------------------------------------------
# Helpers accès Learn
# ------------------------------------------------------
def _learn_require_profile(cur, u: dict, id_effectif: str) -> dict:
    eid = (id_effectif or "").strip()
    if not eid:
        raise HTTPException(status_code=400, detail="id_effectif manquant.")

    profile = learn_fetch_profile(
        cur,
        id_effectif=eid,
        email=(u.get("email") or ""),
        is_super_admin=bool(u.get("is_super_admin")),
    )

    oid = (profile.get("id_owner") or "").strip()
    if not oid:
        raise HTTPException(status_code=403, detail="Profil Learn sans owner.")

    role = (profile.get("role_code") or "user").strip().lower()
    if role not in ("admin", "supervisor", "user"):
        role = "user"

    profile["role_code"] = role
    return profile


def _role_rank(role_code: str) -> int:
    c = (role_code or "").strip().lower()
    if c == "admin":
        return 3
    if c == "supervisor":
        return 2
    return 1


def _learn_require_min_role(profile: dict, min_role: str) -> None:
    if _role_rank(profile.get("role_code")) < _role_rank(min_role):
        raise HTTPException(status_code=403, detail="Accès refusé : droits insuffisants.")


def _competence_exists_owner(cur, oid: str, id_comp: str) -> bool:
    cur.execute(
        """
        SELECT 1
        FROM public.tbl_competence
        WHERE id_owner = %s
          AND id_comp = %s
        LIMIT 1
        """,
        (oid, id_comp),
    )
    return cur.fetchone() is not None


def _next_comp_code(cur, oid: str) -> str:
    cur.execute("SELECT pg_advisory_xact_lock(hashtext(%s))", (f"learn_comp_code:{oid}:CO",))

    cur.execute(
        """
        SELECT COALESCE(
          MAX((regexp_match(code, '^CO([0-9]{5})$'))[1]::int),
          0
        ) AS max_n
        FROM public.tbl_competence
        WHERE id_owner = %s
          AND code ~ '^CO[0-9]{5}$'
        """,
        (oid,),
    )

    r = cur.fetchone() or {}
    max_n_raw = r.get("max_n")
    max_n = int(max_n_raw) if max_n_raw is not None else 0
    nxt = max_n + 1

    if nxt > 99999:
        raise HTTPException(status_code=400, detail="Limite de numérotation atteinte (CO99999).")

    return f"CO{nxt:05d}"


def _level_score(txt: Optional[str]) -> int:
    t = (txt or "").lower()
    score = 0

    for w in ("transmet", "forme", "mentor", "optimise", "anticipe", "industrialise", "pilot", "stratég"):
        if w in t:
            score += 3

    for w in ("autonome", "structure", "fiable", "standard", "applique", "met en oeuvre", "met en œuvre", "gère"):
        if w in t:
            score += 1

    for w in ("guid", "supervis", "avec aide", "assist", "simple"):
        if w in t:
            score -= 2

    return score


def _fix_abc_levels(data: dict) -> None:
    a = data.get("niveaua") or ""
    b = data.get("niveaub") or ""
    c = data.get("niveauc") or ""

    sa = _level_score(a)
    sb = _level_score(b)
    sc = _level_score(c)

    if sa > sc:
        levels = [("niveaua", a, sa), ("niveaub", b, sb), ("niveauc", c, sc)]
        levels.sort(key=lambda x: x[2])
        data["niveaua"] = levels[0][1]
        data["niveaub"] = levels[1][1]
        data["niveauc"] = levels[2][1]


def _pdf_latin1_safe(v: Any) -> str:
    return str(v or "").encode("latin-1", "replace").decode("latin-1")


def _pdf_safe_filename_part(v: Any, max_len: int = 120) -> str:
    s = str(v or "").strip()
    if not s:
        return "document"

    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = re.sub(r"[^A-Za-z0-9._ -]+", "", s)
    s = re.sub(r"\s+", " ", s).strip(" ._-")

    if not s:
        s = "document"

    if len(s) > max_len:
        s = s[:max_len].strip(" ._-")

    return s or "document"


def _fetch_owner_logo_bytes(cur, oid: str) -> Optional[bytes]:
    cur.execute(
        """
        SELECT logo_bytes
        FROM public.tbl_studio_owner_logo
        WHERE id_owner = %s
          AND COALESCE(archive, FALSE) = FALSE
        ORDER BY date_maj DESC, date_creation DESC
        LIMIT 1
        """,
        (oid,),
    )

    row = cur.fetchone() or {}
    raw = row.get("logo_bytes")

    if raw is None:
        return None

    try:
        return bytes(raw)
    except Exception:
        return raw


def _clean_doc_text(value: Any, max_len: int = 18000) -> str:
    txt = str(value or "").replace("\x00", " ")
    txt = re.sub(r"\s+", " ", txt).strip()

    if len(txt) > max_len:
        txt = txt[:max_len].rsplit(" ", 1)[0].strip()

    return txt


async def _extract_document_text(upload: UploadFile) -> str:
    filename = (upload.filename or "").strip()
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    raw = await upload.read()

    if not raw:
        return ""

    if len(raw) > 8 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Document trop volumineux (8 Mo maximum).")

    if ext in ("txt", "md", "csv"):
        return _clean_doc_text(raw.decode("utf-8", errors="ignore"))

    if ext == "pdf":
        if PdfReader is None:
            raise HTTPException(status_code=500, detail="Lecture PDF indisponible côté serveur.")

        try:
            reader = PdfReader(BytesIO(raw))
            parts = []

            for page in reader.pages[:25]:
                try:
                    parts.append(page.extract_text() or "")
                except Exception:
                    pass

            return _clean_doc_text("\n".join(parts))
        except HTTPException:
            raise
        except Exception:
            raise HTTPException(status_code=400, detail="Impossible de lire le document PDF.")

    if ext == "docx":
        if DocxDocument is None:
            raise HTTPException(status_code=500, detail="Lecture DOCX indisponible côté serveur.")

        try:
            doc = DocxDocument(BytesIO(raw))
            parts = [p.text for p in doc.paragraphs if p.text]

            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join((cell.text or "").strip() for cell in row.cells))

            return _clean_doc_text("\n".join(parts))
        except Exception:
            raise HTTPException(status_code=400, detail="Impossible de lire le document DOCX.")

    raise HTTPException(
        status_code=400,
        detail="Format non pris en charge. Formats acceptés : PDF, DOCX, TXT, MD, CSV.",
    )


def _normalize_grille(grille: Any) -> dict:
    if grille is None:
        grille_obj = {}
    elif isinstance(grille, dict):
        grille_obj = grille
    elif isinstance(grille, str):
        try:
            grille_obj = json.loads(grille) if grille.strip() else {}
        except Exception:
            grille_obj = {}
    else:
        grille_obj = {}

    rows = []

    for i in range(1, 5):
        key = f"Critere{i}"
        node = grille_obj.get(key) if isinstance(grille_obj, dict) else {}

        if not isinstance(node, dict):
            node = {}

        ev = node.get("Eval") if isinstance(node.get("Eval"), list) else []

        nom = str(node.get("Nom") or "").strip()
        evals = [
            str(ev[j] if j < len(ev) and ev[j] is not None else "").strip()
            for j in range(4)
        ]

        if not nom:
            continue

        rows.append({
            "Nom": nom,
            "Eval": evals,
        })

    out = {}

    for i in range(1, 5):
        if i <= len(rows):
            out[f"Critere{i}"] = rows[i - 1]
        else:
            out[f"Critere{i}"] = {
                "Nom": "",
                "Eval": ["", "", "", ""],
            }

    return out


# ------------------------------------------------------
# Models
# ------------------------------------------------------
class CreateCompetencePayload(BaseModel):
    intitule: str
    description: Optional[str] = None
    domaine: Optional[str] = None
    niveaua: Optional[str] = None
    niveaub: Optional[str] = None
    niveauc: Optional[str] = None
    grille_evaluation: Optional[Any] = None
    etat: Optional[str] = None


class UpdateCompetencePayload(BaseModel):
    intitule: Optional[str] = None
    description: Optional[str] = None
    domaine: Optional[str] = None
    niveaua: Optional[str] = None
    niveaub: Optional[str] = None
    niveauc: Optional[str] = None
    grille_evaluation: Optional[Any] = None
    etat: Optional[str] = None


class AiDraftCompetencePayload(BaseModel):
    objectif: str
    contexte: Optional[str] = None
    domaine_id: Optional[str] = None
    nb_criteres: Optional[int] = None
    document_text: Optional[str] = None


# ------------------------------------------------------
# IA
# ------------------------------------------------------
def _build_ai_schema() -> dict:
    crit_schema = {
        "type": "object",
        "additionalProperties": False,
        "required": ["Nom", "Eval"],
        "properties": {
            "Nom": {"type": "string", "maxLength": 140},
            "Eval": {
                "type": "array",
                "minItems": 4,
                "maxItems": 4,
                "items": {"type": "string", "maxLength": 120},
            },
        },
    }

    return {
        "type": "object",
        "additionalProperties": False,
        "required": [
            "intitule",
            "description",
            "niveaua",
            "niveaub",
            "niveauc",
            "domaine_id",
            "grille_evaluation",
        ],
        "properties": {
            "intitule": {"type": "string", "minLength": 1, "maxLength": 140},
            "description": {"type": "string", "maxLength": 1200},
            "niveaua": {"type": "string", "minLength": 40, "maxLength": 230},
            "niveaub": {"type": "string", "minLength": 40, "maxLength": 230},
            "niveauc": {"type": "string", "minLength": 40, "maxLength": 230},
            "domaine_id": {"type": ["string", "null"]},
            "grille_evaluation": {
                "type": "object",
                "additionalProperties": False,
                "required": ["Critere1", "Critere2", "Critere3", "Critere4"],
                "properties": {
                    "Critere1": crit_schema,
                    "Critere2": crit_schema,
                    "Critere3": crit_schema,
                    "Critere4": crit_schema,
                },
            },
        },
    }


def _run_ai_draft(cur, oid: str, payload: AiDraftCompetencePayload) -> dict:
    if OpenAI is None:
        raise HTTPException(status_code=500, detail="Lib OpenAI manquante (pip install openai).")

    objectif = (payload.objectif or "").strip()
    if not objectif:
        raise HTTPException(status_code=400, detail="Objectif obligatoire.")

    contexte = (payload.contexte or "").strip()
    domaine_force = (payload.domaine_id or "").strip()
    document_text = _clean_doc_text(payload.document_text or "")

    nb = payload.nb_criteres if payload.nb_criteres is not None else 3
    try:
        nb = int(nb)
    except Exception:
        nb = 3

    if nb not in (1, 2, 3, 4):
        nb = 3

    model = (os.getenv("OPENAI_MODEL_COMP_DRAFT") or "gpt-4o-mini").strip()
    api_key = (os.getenv("OPENAI_API_KEY") or "").strip()

    if not api_key:
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY non configurée.")

    cur.execute(
        """
        SELECT
          id_domaine_competence,
          COALESCE(titre_court, titre, id_domaine_competence) AS label
        FROM public.tbl_domaine_competence
        WHERE COALESCE(masque, FALSE) = FALSE
        ORDER BY
          COALESCE(ordre_affichage, 999999),
          lower(COALESCE(titre_court, titre, id_domaine_competence))
        """
    )

    dom_rows = cur.fetchall() or []
    dom_map = {
        (r.get("id_domaine_competence") or "").strip(): (r.get("label") or "").strip()
        for r in dom_rows
    }

    dom_list_txt = "\n".join([f"- {k} : {v}" for k, v in dom_map.items()]) if dom_map else "- (aucun domaine)"

    sys = (
        "Tu es concepteur pédagogique et tu aides à concevoir une fiche compétence opérationnelle pour Novoskill Learn. "
        "Tu dois respecter STRICTEMENT le schéma JSON fourni. "

        "Avant de rédiger, tu dois décider si la compétence est transverse, métier réutilisable, ou spécifique. "
        "Compétence transverse : elle peut être réutilisée dans plusieurs métiers ou domaines. Son titre, sa description, "
        "ses niveaux A/B/C et ses critères ne doivent pas citer un secteur, un produit, une réglementation, un outil ou une formation précise. "
        "Compétence métier réutilisable : elle appartient à un domaine métier, mais reste exploitable dans plusieurs formations du même domaine. "
        "Son titre peut contenir le domaine métier, mais pas un cas trop étroit. "
        "Compétence spécifique : elle dépend fortement d'une formation, d'un produit, d'une réglementation, d'un dispositif, "
        "d'un outil, d'un secteur ou d'un contexte précis. Dans ce cas, cette spécificité doit apparaître clairement dans l'intitulé. "

        "Règle impérative de cohérence catalogue : il est interdit de produire un intitulé générique avec une description, "
        "des niveaux ou des critères spécifiques. Si le contenu est spécifique, l'intitulé doit être spécifique. "
        "Inversement, si l'intitulé est générique, toute la fiche doit rester générique et réutilisable. "
        "Ne maquille jamais une compétence spécifique avec un titre transverse. "

        "Exemples de mauvais intitulés si le contenu parle de VEFA, PSLA, immobilier neuf ou contrat de réservation : "
        "'Présenter une offre commerciale', 'Conclure une vente', 'Traiter le financement'. "
        "Dans ce cas, l'intitulé doit assumer le contexte, par exemple : "
        "'Présenter une offre immobilière en VEFA', 'Conclure une réservation en immobilier neuf', "
        "'Traiter le financement d'une acquisition en PSLA'. "

        "La compétence doit rester générique et réutilisable uniquement si le contexte précis n'est pas nécessaire "
        "pour comprendre, exercer et évaluer la compétence. "
        "Règles de niveau A/B/C : A = initial guidé, B = avancé autonome, C = expert qui optimise/transmet. "
        "Ne mets jamais un simple label type Initial/Avancé/Expert dans niveaua/b/c. "
        "Rédige des attendus concrets, observables, orientés action. "
        "Les critères d'évaluation doivent être distincts, progressifs, observables et utilisables en formation. "
        "Chaque évaluation doit être courte, 1 phrase, verbe d'action + résultat observable. "
        "Si un document est fourni, utilise-le comme source de contexte, sans recopier de longs passages."
    )

    user = (
        f"Objectif: {objectif}\n"
        f"Contexte: {contexte}\n"
        f"Domaine imposé (si non vide): {domaine_force}\n"
        f"Nombre de critères à produire: {nb}\n\n"
        f"Domaines disponibles (id -> titre_court):\n{dom_list_txt}\n\n"
        f"Document analysé (si présent, extrait nettoyé):\n{document_text}\n\n"
        "Contraintes:\n"
        "- Analyse d'abord mentalement la réutilisabilité de la compétence : transverse, métier réutilisable ou spécifique.\n"
        "- Si la compétence peut être transverse, rédige une fiche entièrement transverse : aucun exemple, critère ou niveau ne doit dépendre du contexte de la formation.\n"
        "- Si la compétence dépend d'un contexte métier précis, rends ce domaine visible dans l'intitulé.\n"
        "- Si la compétence dépend d'un dispositif, produit, réglementation, outil ou cas d'usage précis, cet élément doit apparaître dans l'intitulé.\n"
        "- Interdiction absolue : intitulé générique + description spécifique.\n"
        "- Interdiction absolue : intitulé générique + critères faisant référence au contexte de formation.\n"
        "- Ne crée pas une compétence trop large si l'objectif fourni est visiblement ciblé sur un contexte métier précis.\n"
        "- Ne crée pas une compétence trop spécifique si le même savoir-faire peut être utile dans plusieurs domaines sans perdre son sens.\n"
        f"- Produis exactement {nb} critères NON VIDES (Critere1..Critere{nb}).\n"
        f"- Critere{nb + 1}..Critere4 doivent être VIDES (Nom=\"\" + 4 Eval vides).\n"
        "- Les 4 niveaux d’un critère doivent montrer une progression claire : guidé → autonome → optimisation → expertise/transmission.\n"
        "- Niveaux A/B/C <=230 caractères chacun.\n"
    )

    client = OpenAI(api_key=api_key)

    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": sys},
            {"role": "user", "content": user},
        ],
        temperature=0.3,
        max_tokens=1400,
        response_format={
            "type": "json_schema",
            "json_schema": {
                "name": "learn_competence_draft",
                "schema": _build_ai_schema(),
                "strict": True,
            },
        },
    )

    content = (resp.choices[0].message.content or "").strip()
    if not content:
        raise HTTPException(status_code=500, detail="Réponse IA vide.")

    data = json.loads(content)

    def _bad_level(s: str) -> bool:
        t = (s or "").strip().lower()
        return t in ("initial", "avancé", "avance", "expert", "débutant", "debutant") or len(t) < 40

    if (
        _bad_level(data.get("niveaua", ""))
        or _bad_level(data.get("niveaub", ""))
        or _bad_level(data.get("niveauc", ""))
    ):
        raise HTTPException(
            status_code=400,
            detail="IA: niveaux A/B/C trop courts ou réduits à un label. Regénère avec plus de contexte.",
        )

    _fix_abc_levels(data)

    dom_out = (domaine_force or (data.get("domaine_id") or "")).strip() or None
    if dom_out and dom_out not in dom_map:
        dom_out = None

    data["domaine_id"] = dom_out
    data["grille_evaluation"] = _normalize_grille(data.get("grille_evaluation"))

    used = 0
    for k in ("Critere1", "Critere2", "Critere3", "Critere4"):
        if ((data["grille_evaluation"].get(k) or {}).get("Nom") or "").strip():
            used += 1

    if used < 1:
        raise HTTPException(status_code=400, detail="IA: aucun critère exploitable généré.")

    return data

def _mark_lms_publications_outdated_for_competence(cur, oid: str, id_comp: str) -> int:
    cid = str(id_comp or "").strip()
    if not cid:
        return 0

    comp_json = Json([cid])

    cur.execute(
        """
        WITH formations_touchees AS (
            SELECT DISTINCT ff.id_form
            FROM public.tbl_fiche_formation ff
            WHERE ff.id_owner = %s
              AND COALESCE(ff.archive, FALSE) = FALSE
              AND COALESCE(ff.masque, FALSE) = FALSE
              AND (
                    COALESCE(ff.competences_stagiaires, '[]'::jsonb) @> %s::jsonb
                 OR COALESCE(ff.competences_formateurs, '[]'::jsonb) @> %s::jsonb
                 OR EXISTS (
                    SELECT 1
                    FROM public.tbl_contenu_ligne l
                    WHERE l.id_owner = ff.id_owner
                      AND l.id_form = ff.id_form
                      AND COALESCE(l.archive, FALSE) = FALSE
                      AND (
                            l.id_competence = %s
                         OR COALESCE(l.competences_liees, '[]'::jsonb) @> %s::jsonb
                      )
                 )
              )
        )
        UPDATE public.tbl_learn_lms_publication lp
        SET sync_status = 'outdated',
            updated_at = NOW()
        FROM formations_touchees ft
        WHERE lp.id_owner = %s
          AND lp.id_form = ft.id_form
          AND COALESCE(lp.archive, FALSE) = FALSE
          AND lp.external_id IS NOT NULL
          AND COALESCE(lp.sync_status, '') IN ('synced', 'linked', 'outdated')
        RETURNING lp.id_publication
        """,
        (oid, comp_json, comp_json, cid, comp_json, oid),
    )

    return len(cur.fetchall() or [])

# ------------------------------------------------------
# Endpoints
# ------------------------------------------------------
@router.get("/learn/competences/{id_effectif}/context")
def learn_competences_context(id_effectif: str, request: Request):
    auth = request.headers.get("Authorization", "")
    u = learn_require_user(auth)

    with get_conn() as conn:
        with conn.cursor(row_factory=dict_row) as cur:
            profile = _learn_require_profile(cur, u, id_effectif)

    return {
        "id_owner": profile.get("id_owner"),
        "id_effectif": profile.get("id_effectif"),
        "role_code": profile.get("role_code"),
        "role_label": profile.get("role_label"),
        "can_edit": _role_rank(profile.get("role_code")) >= 2,
    }


@router.get("/learn/competences/{id_effectif}/domaines")
def learn_competences_domaines(id_effectif: str, request: Request):
    auth = request.headers.get("Authorization", "")
    u = learn_require_user(auth)

    try:
        outdated_publications = 0

        with get_conn() as conn:
            with conn.cursor(row_factory=dict_row) as cur:
                _learn_require_profile(cur, u, id_effectif)

                cur.execute(
                    """
                    SELECT
                      id_domaine_competence,
                      titre,
                      titre_court,
                      description,
                      ordre_affichage,
                      couleur
                    FROM public.tbl_domaine_competence
                    WHERE COALESCE(masque, FALSE) = FALSE
                    ORDER BY
                      COALESCE(ordre_affichage, 999999),
                      lower(COALESCE(titre_court, titre, id_domaine_competence))
                    """
                )

                rows = cur.fetchall() or []

        return {"items": rows}

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"learn/competences/domaines error: {e}")


@router.get("/learn/competences/{id_effectif}")
def learn_competences_list(
    id_effectif: str,
    request: Request,
    q: str = "",
    show: str = "active",
):
    auth = request.headers.get("Authorization", "")
    u = learn_require_user(auth)

    try:
        qq = (q or "").strip()
        sh = (show or "active").strip().lower()

        with get_conn() as conn:
            with conn.cursor(row_factory=dict_row) as cur:
                profile = _learn_require_profile(cur, u, id_effectif)
                oid = (profile.get("id_owner") or "").strip()

                where = ["c.id_owner = %s"]
                params = [oid]

                if sh == "active":
                    where.append("COALESCE(c.masque, FALSE) = FALSE")
                    where.append("COALESCE(c.etat, 'active') IN ('active', 'valide')")
                elif sh == "validation":
                    where.append("COALESCE(c.masque, FALSE) = FALSE")
                    where.append("COALESCE(c.etat, '') = 'à valider'")
                elif sh == "archived":
                    where.append("COALESCE(c.masque, FALSE) = TRUE")
                else:
                    where.append("COALESCE(c.masque, FALSE) = FALSE")

                if qq:
                    like = f"%{qq}%"
                    where.append(
                        """
                        (
                          c.code ILIKE %s
                          OR c.intitule ILIKE %s
                          OR COALESCE(c.description, '') ILIKE %s
                          OR COALESCE(dc.titre_court, '') ILIKE %s
                          OR COALESCE(dc.titre, '') ILIKE %s
                        )
                        """
                    )
                    params.extend([like, like, like, like, like])

                cur.execute(
                    f"""
                    SELECT
                      c.id_comp,
                      c.code,
                      c.intitule,
                      c.domaine,
                      dc.titre_court AS domaine_titre_court,
                      dc.titre AS domaine_titre,
                      dc.couleur AS domaine_couleur,
                      c.etat,
                      COALESCE(c.masque, FALSE) AS masque,
                      c.date_modification
                    FROM public.tbl_competence c
                    LEFT JOIN public.tbl_domaine_competence dc
                      ON dc.id_domaine_competence = c.domaine
                     AND COALESCE(dc.masque, FALSE) = FALSE
                    WHERE {" AND ".join(where)}
                    ORDER BY
                      lower(COALESCE(c.code, '')),
                      lower(COALESCE(c.intitule, ''))
                    """,
                    tuple(params),
                )

                rows = cur.fetchall() or []

        return {"items": rows}

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"learn/competences list error: {e}")


@router.post("/learn/competences/{id_effectif}")
def learn_competence_create(id_effectif: str, payload: CreateCompetencePayload, request: Request):
    auth = request.headers.get("Authorization", "")
    u = learn_require_user(auth)

    try:
        title = (payload.intitule or "").strip()
        if not title:
            raise HTTPException(status_code=400, detail="Intitulé obligatoire.")

        with get_conn() as conn:
            with conn.cursor(row_factory=dict_row) as cur:
                profile = _learn_require_profile(cur, u, id_effectif)
                _learn_require_min_role(profile, "supervisor")
                oid = (profile.get("id_owner") or "").strip()

                code = _next_comp_code(cur, oid)
                cid = str(uuid.uuid4())

                cur.execute(
                    """
                    INSERT INTO public.tbl_competence
                      (id_comp, id_owner, code, intitule, description, domaine,
                       niveaua, niveaub, niveauc, grille_evaluation,
                       etat, masque, date_creation, date_modification)
                    VALUES
                      (%s, %s, %s, %s, %s, %s,
                       %s, %s, %s, %s,
                       %s, FALSE, NOW(), NOW())
                    """,
                    (
                        cid,
                        oid,
                        code,
                        title,
                        (payload.description or None),
                        (payload.domaine or None),
                        (payload.niveaua or None),
                        (payload.niveaub or None),
                        (payload.niveauc or None),
                        Json(_normalize_grille(payload.grille_evaluation)),
                        (payload.etat or "à valider"),
                    ),
                )

                conn.commit()

        return {"id_comp": cid, "code": code}

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"learn/competences create error: {e}")


@router.post("/learn/competences/{id_effectif}/draft/ai")
def learn_competence_ai_draft(id_effectif: str, payload: AiDraftCompetencePayload, request: Request):
    auth = request.headers.get("Authorization", "")
    u = learn_require_user(auth)

    try:
        with get_conn() as conn:
            with conn.cursor(row_factory=dict_row) as cur:
                profile = _learn_require_profile(cur, u, id_effectif)
                _learn_require_min_role(profile, "supervisor")
                oid = (profile.get("id_owner") or "").strip()
                return _run_ai_draft(cur, oid, payload)

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"learn/competences ai_draft error: {e}")


@router.post("/learn/competences/{id_effectif}/draft/ai-document")
async def learn_competence_ai_draft_document(
    id_effectif: str,
    request: Request,
    objectif: str = Form(...),
    contexte: Optional[str] = Form(None),
    domaine_id: Optional[str] = Form(None),
    nb_criteres: Optional[int] = Form(3),
    document: UploadFile = File(...),
):
    auth = request.headers.get("Authorization", "")
    u = learn_require_user(auth)

    try:
        doc_text = await _extract_document_text(document)

        payload = AiDraftCompetencePayload(
            objectif=objectif,
            contexte=contexte,
            domaine_id=domaine_id,
            nb_criteres=nb_criteres,
            document_text=doc_text,
        )

        with get_conn() as conn:
            with conn.cursor(row_factory=dict_row) as cur:
                profile = _learn_require_profile(cur, u, id_effectif)
                _learn_require_min_role(profile, "supervisor")
                oid = (profile.get("id_owner") or "").strip()
                return _run_ai_draft(cur, oid, payload)

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"learn/competences ai_draft_document error: {e}")


@router.get("/learn/competences/{id_effectif}/{id_comp}")
def learn_competence_detail(id_effectif: str, id_comp: str, request: Request):
    auth = request.headers.get("Authorization", "")
    u = learn_require_user(auth)

    try:
        cid = (id_comp or "").strip()
        if not cid:
            raise HTTPException(status_code=400, detail="id_comp manquant.")

        with get_conn() as conn:
            with conn.cursor(row_factory=dict_row) as cur:
                profile = _learn_require_profile(cur, u, id_effectif)
                oid = (profile.get("id_owner") or "").strip()

                cur.execute(
                    """
                    SELECT
                      c.id_comp,
                      c.code,
                      c.intitule,
                      c.description,
                      c.domaine,
                      dc.titre_court AS domaine_titre_court,
                      dc.titre AS domaine_titre,
                      dc.couleur AS domaine_couleur,
                      c.niveaua,
                      c.niveaub,
                      c.niveauc,
                      c.grille_evaluation,
                      c.etat,
                      COALESCE(c.masque, FALSE) AS masque,
                      c.date_creation,
                      c.date_modification
                    FROM public.tbl_competence c
                    LEFT JOIN public.tbl_domaine_competence dc
                      ON dc.id_domaine_competence = c.domaine
                     AND COALESCE(dc.masque, FALSE) = FALSE
                    WHERE c.id_owner = %s
                      AND c.id_comp = %s
                    LIMIT 1
                    """,
                    (oid, cid),
                )

                r = cur.fetchone()
                if not r:
                    raise HTTPException(status_code=404, detail="Compétence introuvable.")

        r["grille_evaluation"] = _normalize_grille(r.get("grille_evaluation"))
        return r

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"learn/competences detail error: {e}")


@router.post("/learn/competences/{id_effectif}/{id_comp}")
def learn_competence_update(id_effectif: str, id_comp: str, payload: UpdateCompetencePayload, request: Request):
    auth = request.headers.get("Authorization", "")
    u = learn_require_user(auth)

    try:
        cid = (id_comp or "").strip()
        if not cid:
            raise HTTPException(status_code=400, detail="id_comp manquant.")

        patch_fields = payload.__fields_set__ or set()
        if not patch_fields:
            return {"ok": True}

        with get_conn() as conn:
            with conn.cursor(row_factory=dict_row) as cur:
                profile = _learn_require_profile(cur, u, id_effectif)
                _learn_require_min_role(profile, "supervisor")
                oid = (profile.get("id_owner") or "").strip()

                if not _competence_exists_owner(cur, oid, cid):
                    raise HTTPException(status_code=404, detail="Compétence introuvable.")

                cols = []
                vals = []

                if "intitule" in patch_fields:
                    title = (payload.intitule or "").strip()
                    if not title:
                        raise HTTPException(status_code=400, detail="Intitulé obligatoire.")
                    cols.append("intitule = %s")
                    vals.append(title)

                if "description" in patch_fields:
                    cols.append("description = %s")
                    vals.append(payload.description)

                if "domaine" in patch_fields:
                    cols.append("domaine = %s")
                    vals.append(payload.domaine)

                if "niveaua" in patch_fields:
                    cols.append("niveaua = %s")
                    vals.append(payload.niveaua)

                if "niveaub" in patch_fields:
                    cols.append("niveaub = %s")
                    vals.append(payload.niveaub)

                if "niveauc" in patch_fields:
                    cols.append("niveauc = %s")
                    vals.append(payload.niveauc)

                if "grille_evaluation" in patch_fields:
                    cols.append("grille_evaluation = %s")
                    vals.append(Json(_normalize_grille(payload.grille_evaluation)))

                if "etat" in patch_fields:
                    etat = (payload.etat or "").strip() or "à valider"
                    cols.append("etat = %s")
                    vals.append(etat)

                if cols:
                    cols.append("date_modification = NOW()")
                    vals.extend([cid, oid])

                    cur.execute(
                        f"""
                        UPDATE public.tbl_competence
                        SET {", ".join(cols)}
                        WHERE id_comp = %s
                          AND id_owner = %s
                        """,
                        tuple(vals),
                    )

                    outdated_publications = _mark_lms_publications_outdated_for_competence(cur, oid, cid)

                    conn.commit()

        return {
            "ok": True,
            "lms_publications_outdated": outdated_publications,
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"learn/competences update error: {e}")


@router.post("/learn/competences/{id_effectif}/{id_comp}/archive")
def learn_competence_archive(id_effectif: str, id_comp: str, request: Request):
    auth = request.headers.get("Authorization", "")
    u = learn_require_user(auth)

    try:
        cid = (id_comp or "").strip()
        if not cid:
            raise HTTPException(status_code=400, detail="id_comp manquant.")

        with get_conn() as conn:
            with conn.cursor(row_factory=dict_row) as cur:
                profile = _learn_require_profile(cur, u, id_effectif)
                _learn_require_min_role(profile, "supervisor")
                oid = (profile.get("id_owner") or "").strip()

                cur.execute(
                    """
                    UPDATE public.tbl_competence
                    SET masque = TRUE,
                        date_modification = NOW()
                    WHERE id_owner = %s
                      AND id_comp = %s
                      AND COALESCE(masque, FALSE) = FALSE
                    """,
                    (oid, cid),
                )

                conn.commit()

        return {"ok": True}

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"learn/competences archive error: {e}")


@router.get("/learn/competences/{id_effectif}/{id_comp}/fiche_pdf")
def learn_competence_fiche_pdf(id_effectif: str, id_comp: str, request: Request):
    auth = request.headers.get("Authorization", "")
    u = learn_require_user(auth)

    try:
        cid = (id_comp or "").strip()
        if not cid:
            raise HTTPException(status_code=400, detail="id_comp manquant.")

        with get_conn() as conn:
            with conn.cursor(row_factory=dict_row) as cur:
                profile = _learn_require_profile(cur, u, id_effectif)
                oid = (profile.get("id_owner") or "").strip()

                cur.execute(
                    """
                    SELECT
                      c.id_comp,
                      c.code,
                      c.intitule,
                      c.description,
                      c.domaine,
                      c.niveaua,
                      c.niveaub,
                      c.niveauc,
                      c.grille_evaluation,
                      dc.titre_court AS domaine_titre_court,
                      dc.titre AS domaine_titre
                    FROM public.tbl_competence c
                    LEFT JOIN public.tbl_domaine_competence dc
                      ON dc.id_domaine_competence = c.domaine
                     AND COALESCE(dc.masque, FALSE) = FALSE
                    WHERE c.id_owner = %s
                      AND c.id_comp = %s
                      AND COALESCE(c.masque, FALSE) = FALSE
                      AND COALESCE(c.etat, 'active') IN ('active', 'valide', 'à valider')
                    LIMIT 1
                    """,
                    (oid, cid),
                )

                row = cur.fetchone() or {}
                if not row:
                    raise HTTPException(status_code=404, detail="Compétence introuvable dans le référentiel courant.")

                logo_bytes = _fetch_owner_logo_bytes(cur, oid)

        skill = {
            "id_comp": row.get("id_comp"),
            "code": (row.get("code") or "").strip(),
            "intitule": (row.get("intitule") or "").strip(),
            "description": row.get("description") or "",
            "niveaua": row.get("niveaua") or "",
            "niveaub": row.get("niveaub") or "",
            "niveauc": row.get("niveauc") or "",
            "grille_evaluation": row.get("grille_evaluation"),
            "domaine": row.get("domaine") or "",
            "domaine_titre": (
                (row.get("domaine_titre_court") or "").strip()
                or (row.get("domaine_titre") or "").strip()
            ),
        }

        code_label = skill.get("code") or "Compétence"
        intitule_label = skill.get("intitule") or "Compétence"
        owner_label = (profile.get("nom_owner") or "Novoskill Learn").strip() or "Novoskill Learn"

        filename = _pdf_latin1_safe(
            f"Fiche compétence {_pdf_safe_filename_part(code_label, 32)} - {_pdf_safe_filename_part(intitule_label, 80)}.pdf"
        )

        pdf_bytes = build_pdf_document(
            build_competence_pdf_story(skill),
            meta={
                "title": _pdf_latin1_safe(f"Fiche compétence - {code_label} - {intitule_label}"),
                "doc_label": _pdf_latin1_safe("Fiche compétence"),
                "footer_left": _pdf_latin1_safe("Novoskill Learn • Fiche compétence"),
                "header_right": _pdf_latin1_safe(owner_label),
                "header_right_font_name": "Helvetica-Bold",
                "header_right_font_size": 10.5,
                "logo_bytes": logo_bytes,
            },
        )

        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'inline; filename="{filename}"',
                "Cache-Control": "no-store",
            },
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"learn/competences fiche_pdf error: {e}")